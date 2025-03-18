from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
from tqdm import tqdm

class JsonToWordConverter:
    def __init__(self):
        self.document = None
        
    def create_heading(self, text, level=1):
        heading = self.document.add_heading('', level=level)
        run = heading.add_run(text)
        run.font.size = Pt(14 - level)
        run.font.color.rgb = RGBColor(0, 0, 0)
        
    def add_category(self, category):
        # Add main category
        self.create_heading(f"{category['code']} - {category['name']}", 1)
        
        # Add subcategories
        for subcategory in category.get('subcategories', []):
            self.create_heading(f"{subcategory['code']} - {subcategory['name']}", 2)
            
            # Add sub-subcategories
            for sub_sub in subcategory.get('subSubcategories', []):
                self.create_heading(f"{sub_sub['code']} - {sub_sub['name']}", 3)
                
        self.document.add_paragraph()  # Add spacing between categories
        
    def convert(self, data, output_file):
        self.document = Document()
        
        # Add title
        title = self.document.add_heading('Machine Categories', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Process each category
        for category in data['categories']:
            if category:  # Check if category data exists
                self.add_category(category)
        
        # Save the document
        self.document.save(output_file)

    def format_heading(self, level):
        """Format heading based on level"""
        if level == 0:
            return self.document.styles['Heading 1']
        elif level == 1:
            return self.document.styles['Heading 2']
        else:
            return self.document.styles['Heading 3']

    def add_item(self, key, value, level=0):
        """Add item to document with proper indentation and formatting"""
        # Add indentation based on level
        indent = "    " * level
        
        if isinstance(value, dict):
            # Add category/key as heading
            heading = self.document.add_paragraph(style=self.format_heading(level))
            heading.add_run(f"{key}").bold = True
            
            # Recursively process nested dictionary
            for sub_key, sub_value in value.items():
                self.add_item(sub_key, sub_value, level + 1)
                
        elif isinstance(value, list):
            # Add category/key as heading
            heading = self.document.add_paragraph(style=self.format_heading(level))
            heading.add_run(f"{key}").bold = True
            
            # Process list items
            for idx, item in enumerate(value):
                if isinstance(item, dict):
                    for sub_key, sub_value in item.items():
                        self.add_item(sub_key, sub_value, level + 1)
                else:
                    p = self.document.add_paragraph()
                    p.paragraph_format.left_indent = Inches(level * 0.5)
                    p.add_run(f"• {str(item)}")
        else:
            # Add key-value pair as regular text
            p = self.document.add_paragraph()
            p.paragraph_format.left_indent = Inches(level * 0.5)
            p.add_run(f"{key}: ").bold = True
            p.add_run(str(value))

    def process_list(self, data_list, level=0):
        """Process a list of items"""
        for idx, item in enumerate(tqdm(data_list, desc="Converting")):
            if isinstance(item, dict):
                for key, value in item.items():
                    self.add_item(key, value, level)
            else:
                p = self.document.add_paragraph()
                p.paragraph_format.left_indent = Inches(level * 0.5)
                p.add_run(f"• {str(item)}")

    def process_dict(self, data_dict, level=0):
        """Process a dictionary"""
        for key, value in tqdm(data_dict.items(), desc="Converting"):
            self.add_item(key, value, level)

    def convert_json(self, input_json_path, output_docx_path):
        """Convert JSON file to Word document"""
        # Read JSON file
        with open(input_json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Add title
        title = self.document.add_heading('JSON Structure', 0)
        
        # Process JSON data based on its type
        if isinstance(data, list):
            self.process_list(data)
        elif isinstance(data, dict):
            self.process_dict(data)
        
        # Save document
        self.document.save(output_docx_path)
        print(f"Document saved successfully to {output_docx_path}") 